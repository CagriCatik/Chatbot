import os
import tempfile
import shutil
from typing import Optional, Dict, Union, List
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import UnstructuredPDFLoader
from langchain.schema import Document  # Import Document class
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
import streamlit as st
import logging
from config import config, PERSIST_DIRECTORY

logger = logging.getLogger(__name__)

def extract_text_from_docx(file) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_html(file) -> str:
    """Extract text from an HTML file."""
    content = file.read().decode("utf-8")
    soup = BeautifulSoup(content, "html.parser")
    return soup.get_text()

def create_vector_db(file_upload: Union[st.runtime.uploaded_file_manager.UploadedFile, Dict[str, str]]) -> Optional[Chroma]:
    """
    Create a vector database from an uploaded file (PDF, DOCX, HTML).
    """
    if isinstance(file_upload, dict):  # Handling raw text input
        file_name = file_upload["name"]
        file_text = file_upload["text"]
    elif isinstance(file_upload, st.runtime.uploaded_file_manager.UploadedFile):  # Handling uploaded file
        file_name = file_upload.name
        logger.info(f"Creating vector DB from file upload: {file_name}")

        temp_dir = tempfile.mkdtemp()
        path = os.path.join(temp_dir, file_name)

        try:
            file_type = file_upload.type

            with open(path, "wb") as f:
                f.write(file_upload.getvalue())
                logger.info(f"File saved to temporary path: {path}")

            if file_type == "application/pdf":
                loader = UnstructuredPDFLoader(path)
                data = loader.load()
                file_text = "\n".join([doc.page_content for doc in data])
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_text = extract_text_from_docx(file_upload)
            elif file_type == "text/html":
                file_text = extract_text_from_html(file_upload)
            else:
                st.error(f"Unsupported file type: {file_type}")
                logger.error(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(f"Error processing file {file_name}: {e}")
            st.error(f"Error processing file: {file_name}")
            return None
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)
            logger.info(f"Temporary directory {temp_dir} removed")
    else:
        st.error("Invalid file type passed to `create_vector_db`.")
        return None

    # Convert text into LangChain Document objects
    documents: List[Document] = [Document(page_content=file_text, metadata={"source": file_name})]

    # Use text splitter parameters from config
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=config["text_splitter"]["chunk_size"],
        chunk_overlap=config["text_splitter"]["chunk_overlap"]
    )
    chunks = text_splitter.split_documents(documents)
    logger.info("Document split into chunks")

    # Use embedding model from config
    embeddings = OllamaEmbeddings(model=config["embeddings"]["model"])
    vector_db = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=PERSIST_DIRECTORY,
        collection_name=f"file_{hash(file_name)}"
    )
    logger.info("Vector DB created with persistent storage")

    return vector_db

def delete_vector_db(vector_db: Optional[Chroma]) -> None:
    """
    Delete the vector database and clear related session state.
    """
    logger.info("Deleting vector DB")
    if vector_db is not None:
        try:
            vector_db.delete_collection()
            st.session_state.pop("pdf_pages", None)
            st.session_state.pop("file_upload", None)
            st.session_state.pop("vector_db", None)
            st.success("Collection and temporary files deleted successfully.")
            logger.info("Vector DB and related session state cleared")
            st.rerun()
        except Exception as e:
            st.error(f"Error deleting collection: {str(e)}")
            logger.error(f"Error deleting collection: {e}")
    else:
        st.error("No vector database found to delete.")
        logger.warning("Attempted to delete vector DB, but none was found")
